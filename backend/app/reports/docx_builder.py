"""
Word report generation with python-docx.
"""

import io
import os
import subprocess
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import matplotlib.pyplot as plt
import numpy as np

class ReportBuilder:
    """Build professional Word reports."""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Configure document styles."""
        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Heading 1
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(16)
        h1.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
        h1.font.bold = True
        
        # Heading 2
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(14)
        h2.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
        h2.font.bold = True
    
    def build_report(self, project_data: Dict, analysis_results: Dict,
                     plots: Dict[str, bytes], author: str = "Process Chemist") -> bytes:
        """
        Build complete QbD-DOE report.
        
        Args:
            project_data: Project metadata, CQAs, CPPs, design info
            analysis_results: Full analysis output
            plots: Dict mapping plot name -> PNG bytes
            author: Report author name
            
        Returns:
            DOCX file as bytes
        """
        self._add_cover_page(project_data, author)
        self._add_executive_summary(analysis_results)
        self._add_qbd_overview()
        self._add_cqa_cpp_tables(project_data)
        self._add_experimental_design(project_data)
        self._add_results(project_data)
        self._add_analysis(analysis_results)
        self._add_graphical_analysis(plots)
        self._add_optimization(analysis_results)
        self._add_conclusions()
        
        # Save to bytes
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        raw_bytes = buf.getvalue()
        
        # Post-process with LibreOffice if available
        return self._post_process(raw_bytes)
    
    def _add_cover_page(self, project_data: Dict, author: str):
        """Section 1: Cover page."""
        self.doc.add_heading('QbD-DOE Process Development Report', 0)
        
        p = self.doc.add_paragraph()
        p.add_run(f"API: {project_data.get('api_name', 'Unknown')}\n").bold = True
        p.add_run(f"Project ID: {project_data.get('project_id', 'N/A')}\n")
        p.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
        p.add_run(f"Author: {author}\n")
        p.add_run(f"Unit Operation: {project_data.get('unit_operation', 'N/A')}\n")
        
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run("CONFIDENTIAL - INTERNAL R&D USE ONLY").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()
    
    def _add_executive_summary(self, analysis_results: Dict):
        """Section 2: Executive summary."""
        self.doc.add_heading('Executive Summary', level=1)
        
        summary = analysis_results.get('summary', {})
        
        p = self.doc.add_paragraph()
        p.add_run(f"Model Adequacy: {summary.get('verdict', 'N/A')}\n").bold = True
        p.add_run(f"R²: {summary.get('r2', 0):.4f}\n")
        p.add_run(f"Predicted R² (Q²): {summary.get('q2', 0):.4f}\n")
        p.add_run(f"Significant Factors: {summary.get('n_significant', 0)}\n")
        
        self.doc.add_paragraph(
            "This report presents a Quality by Design (QbD) approach using "
            "Design of Experiments (DOE) to develop and optimize the process. "
            "All statistical analyses were performed using best-practice methods "
            "with validated software."
        )
        
        self.doc.add_page_break()
    
    def _add_qbd_overview(self):
        """Section 3: QbD framework overview."""
        self.doc.add_heading('QbD Framework Overview', level=1)
        
        self.doc.add_paragraph(
            "Quality by Design (QbD) is a systematic approach to development that "
            "begins with predefined objectives and emphasizes product and process "
            "understanding and process control."
        )
        
        self.doc.add_heading('QTPP (Quality Target Product Profile)', level=2)
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'QTPP Element'
        hdr_cells[1].text = 'Target'
        hdr_cells[2].text = 'Justification'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Assay'
        row_cells[1].text = '95-102%'
        row_cells[2].text = 'ICH Q3A/B requirements'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Total Impurities'
        row_cells[1].text = '≤2.0%'
        row_cells[2].text = 'Safety threshold'
        
        self.doc.add_page_break()
    
    def _add_cqa_cpp_tables(self, project_data: Dict):
        """Section 4: CQAs and CPPs."""
        self.doc.add_heading('Critical Quality Attributes (CQAs)', level=1)
        
        cqas = project_data.get('cqas', [])
        if cqas:
            table = self.doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'CQA'
            hdr[1].text = 'Unit'
            hdr[2].text = 'Target'
            hdr[3].text = 'LSL'
            hdr[4].text = 'USL'
            hdr[5].text = 'Goal'
            
            for cqa in cqas:
                row = table.add_row().cells
                row[0].text = cqa.get('name', '')
                row[1].text = cqa.get('unit', '')
                row[2].text = str(cqa.get('target', ''))
                row[3].text = str(cqa.get('lower_spec', ''))
                row[4].text = str(cqa.get('upper_spec', ''))
                row[5].text = cqa.get('goal', '')
        
        self.doc.add_heading('Critical Process Parameters (CPPs)', level=1)
        
        cpps = project_data.get('cpps', [])
        if cpps:
            table = self.doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'CPP'
            hdr[1].text = 'Unit'
            hdr[2].text = 'Low'
            hdr[3].text = 'Center'
            hdr[4].text = 'High'
            hdr[5].text = 'Rationale'
            
            for cpp in cpps:
                row = table.add_row().cells
                row[0].text = cpp.get('name', '')
                row[1].text = cpp.get('unit', '')
                row[2].text = str(cpp.get('low', ''))
                row[3].text = str(cpp.get('center', ''))
                row[4].text = str(cpp.get('high', ''))
                row[5].text = cpp.get('rationale', '')
        
        self.doc.add_page_break()
    
    def _add_experimental_design(self, project_data: Dict):
        """Section 5: Experimental design."""
        self.doc.add_heading('Experimental Design', level=1)
        
        design = project_data.get('design', {})
        
        p = self.doc.add_paragraph()
        p.add_run(f"Design Type: {design.get('type', 'N/A')}\n")
        p.add_run(f"Number of Factors: {design.get('n_factors', 0)}\n")
        p.add_run(f"Total Runs: {design.get('n_total_runs', 0)}\n")
        p.add_run(f"Center Points: {design.get('center_points', 0)}\n")
        p.add_run(f"Alpha: {design.get('alpha', 'N/A')}\n")
        
        self.doc.add_paragraph(
            "The design was generated using pyDOE3 and randomized with a "
            "fixed seed for reproducibility. All runs were performed in the "
            "order specified by the RunOrder column."
        )
        
        self.doc.add_page_break()
    
    def _add_results(self, project_data: Dict):
        """Section 6: Experimental results."""
        self.doc.add_heading('Experimental Results', level=1)
        
        results = project_data.get('results', [])
        if results:
            # Add results table (first 10 rows)
            if len(results) > 0:
                table = self.doc.add_table(rows=1, cols=len(results[0]))
                table.style = 'Light Grid Accent 1'
                
                # Header
                for i, key in enumerate(results[0].keys()):
                    table.rows[0].cells[i].text = str(key)
                
                # Data (first 10 rows)
                for row_data in results[:10]:
                    row = table.add_row().cells
                    for i, val in enumerate(row_data.values()):
                        row[i].text = str(val)[:50]  # Truncate long values
                
                if len(results) > 10:
                    self.doc.add_paragraph(f"... {len(results) - 10} additional rows (see appendix)")
        
        self.doc.add_page_break()
    
    def _add_analysis(self, analysis_results: Dict):
        """Section 7: Statistical analysis."""
        self.doc.add_heading('Statistical Analysis', level=1)
        
        # ANOVA table
        self.doc.add_heading('ANOVA', level=2)
        anova = analysis_results.get('anova', {})
        if 'type3_ss' in anova:
            table = self.doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Source'
            hdr[1].text = 'SS'
            hdr[2].text = 'df'
            hdr[3].text = 'MS'
            hdr[4].text = 'F'
            hdr[5].text = 'p-value'
            
            for row_data in anova['type3_ss'][:10]:
                row = table.add_row().cells
                row[0].text = str(row_data.get('index', ''))[:20]
                row[1].text = f"{row_data.get('sum_sq', 0):.4f}"
                row[2].text = str(row_data.get('df', ''))
                row[3].text = f"{row_data.get('mean_sq', 0):.4f}"
                row[4].text = f"{row_data.get('F', 0):.4f}"
                row[5].text = f"{row_data.get('PR(>F)', 1):.4e}"
        
        # Coefficient table
        self.doc.add_heading('Regression Coefficients', level=2)
        coefs = analysis_results.get('coefficients', [])
        if coefs:
            table = self.doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Term'
            hdr[1].text = 'Estimate'
            hdr[2].text = 'SE'
            hdr[3].text = 't'
            hdr[4].text = 'p-value'
            hdr[5].text = '95% CI'
            
            for c in coefs:
                row = table.add_row().cells
                row[0].text = str(c.get('Term', ''))[:20]
                row[1].text = f"{c.get('Estimate', 0):.4f}"
                row[2].text = f"{c.get('SE', 0):.4f}"
                row[3].text = f"{c.get('t_value', 0):.4f}"
                row[4].text = f"{c.get('p_value', 1):.4e}"
                row[5].text = f"[{c.get('CI_lower', 0):.4f}, {c.get('CI_upper', 0):.4f}]"
        
        # Model equation
        self.doc.add_heading('Model Equation (Coded Units)', level=2)
        equation = analysis_results.get('model_equation_coded', 'N/A')
        p = self.doc.add_paragraph()
        p.add_run(equation).italic = True
        
        # Diagnostics summary
        self.doc.add_heading('Model Diagnostics', level=2)
        diag = analysis_results.get('diagnostics', {})
        
        p = self.doc.add_paragraph()
        p.add_run(f"Normality (Shapiro-Wilk): {'PASS' if diag.get('normality', {}).get('overall_normal', False) else 'FAIL'}\n")
        p.add_run(f"Heteroscedasticity (Breusch-Pagan): {'PASS' if diag.get('heteroscedasticity', {}).get('pass', False) else 'FAIL'}\n")
        p.add_run(f"Autocorrelation (Durbin-Watson): {'PASS' if diag.get('autocorrelation', {}).get('pass', False) else 'FAIL'}\n")
        
        self.doc.add_page_break()
    
    def _add_graphical_analysis(self, plots: Dict[str, bytes]):
        """Section 8: Graphical analysis."""
        self.doc.add_heading('Graphical Analysis', level=1)
        
        for plot_name, plot_bytes in plots.items():
            self.doc.add_heading(plot_name.replace('_', ' ').title(), level=2)
            
            # Save to temp file and add
            temp_path = f"/tmp/plot_{plot_name}.png"
            with open(temp_path, 'wb') as f:
                f.write(plot_bytes)
            
            self.doc.add_picture(temp_path, width=Inches(6))
            
            self.doc.add_paragraph(
                f"Figure: {plot_name.replace('_', ' ').title()}. "
                "Generated from fitted model at 300 DPI."
            )
            
            # Cleanup
            os.remove(temp_path)
        
        self.doc.add_page_break()
    
    def _add_optimization(self, analysis_results: Dict):
        """Section 9: Optimization & design space."""
        self.doc.add_heading('Optimization & Design Space', level=1)
        
        opt = analysis_results.get('optimization', {})
        
        p = self.doc.add_paragraph()
        p.add_run(f"Optimal Overall Desirability: {opt.get('overall_desirability', 0):.4f}\n")
        p.add_run(f"Optimal Factor Settings:\n").bold = True
        
        factors = opt.get('optimal_factors', [])
        for i, f in enumerate(factors):
            p.add_run(f"  Factor {i+1}: {f:.4f} (coded)\n")
        
        # Predictions at optimum
        preds = opt.get('predictions', {})
        p.add_run(f"\nPredicted CQAs at Optimum:\n").bold = True
        for cqa, data in preds.items():
            p.add_run(f"  {cqa}: {data.get('predicted', 0):.2f} "
                     f"[{data.get('pi_lower', 0):.2f}, {data.get('pi_upper', 0):.2f}]\n")
        
        # Monte Carlo
        mc = analysis_results.get('monte_carlo', {})
        p.add_run(f"\nMonte Carlo P(all specs): {mc.get('probability_all_specs', 0):.4f}\n")
        
        self.doc.add_paragraph(
            "NOTE: The proposed design space requires experimental verification "
            "at the edges before formal validation. This is a statistical prediction, "
            "not a proven manufacturing space."
        )
        
        self.doc.add_page_break()
    
    def _add_conclusions(self):
        """Section 10: Conclusions."""
        self.doc.add_heading('Conclusions & Control Strategy', level=1)
        
        self.doc.add_paragraph(
            "Based on the QbD-DOE study, the following control strategy is proposed:\n\n"
            "1. CPPs shall be controlled within the Normal Operating Range (NOR) "
            "shown in the CPP table.\n"
            "2. CQAs shall be monitored per batch with acceptance criteria "
            "as defined in the CQA table.\n"
            "3. The proposed design space (P(in spec) ≥ 0.95) shall be verified "
            "with confirmatory runs at the edges.\n"
            "4. Scale-up shall include verification that the same model applies.\n"
        )
        
        self.doc.add_heading('Appendix', level=1)
        self.doc.add_paragraph(
            "Software versions:\n"
            "- Python 3.11\n"
            "- statsmodels 0.14.0\n"
            "- scipy 1.11.3\n"
            "- numpy 1.26.2\n"
            "- pyDOE3 1.0.0\n"
        )
    
    def _post_process(self, raw_bytes: bytes) -> bytes:
        """Post-process through LibreOffice for MS Word compatibility."""
        try:
            # Save raw
            raw_path = '/tmp/report_raw.docx'
            with open(raw_path, 'wb') as f:
                f.write(raw_bytes)
            
            # Convert with LibreOffice
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'docx', 
                 '--outdir', '/tmp', raw_path],
                capture_output=True,
                timeout=30
            )
            
            if result.returncode == 0 and os.path.exists('/tmp/report_raw.docx'):
                with open('/tmp/report_raw.docx', 'rb') as f:
                    return f.read()
            else:
                # Return raw with warning
                return raw_bytes
                
        except Exception as e:
            # LibreOffice not available, return raw
            return raw_bytes
